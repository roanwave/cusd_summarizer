"""
Generates Word document and text summaries for the CUSD Email Summarizer.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from modules.logger import get_logger

logger = get_logger(__name__)


class DocumentGenerator:
    """Handles Word document generation for the CUSD Email Summarizer."""

    def __init__(self, output_dir: str = "output"):
        """Initialize with optional output directory (kept for compatibility)."""
        self.output_dir = output_dir
        self.doc = Document()

    # ----------------------------------------------------------------------
    # Full compatibility wrapper
    # ----------------------------------------------------------------------
    def create_digest_document(
        self,
        emails=None,
        consolidated_digest=None,
        output_path: str = None,
        *args,
        **kwargs,
    ):
        """
        Maintains backward compatibility with all orchestrator call variants.
        Some older versions pass 'digest', others 'consolidated_digest'.
        """
        digest = consolidated_digest or kwargs.get("digest")
        if digest is None:
            raise ValueError("Missing required digest data.")
        if output_path is None:
            timestamp = datetime.now().strftime("%B_%d_%Y")
            output_path = os.path.join(self.output_dir, f"CUSD_Digest_{timestamp}.docx")

        logger.debug(
            "create_digest_document() invoked "
            f"(emails={len(emails) if emails else 0}, path={output_path})"
        )
        return self.create_document(emails, digest, output_path)

    # ----------------------------------------------------------------------
    def create_document(self, emails, digest, output_path):
        """Build digest and per-email summary sections."""
        try:
            self._add_title()
            self._add_digest(digest)
            self._add_email_summaries(emails)

            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            self.doc.save(output_path)
            logger.info(f"Document saved: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            raise

    # ======================================================================
    # Sections
    # ======================================================================

    def _add_title(self):
        """Add document header."""
        self.doc.add_heading(
            f"CUSD Email Digest - {datetime.now():%B %d, %Y}", level=1
        )
        self.doc.add_paragraph("Compiled by the CUSD Email Summarizer.")
        self.doc.add_paragraph("")

    def _add_digest(self, digest):
        """Add executive summary, events, action items, and announcements."""
        # Ensure digest has all required keys with defaults
        if not isinstance(digest, dict):
            logger.error(f"Invalid digest type: {type(digest)}")
            digest = {}
        
        digest.setdefault('executive_summary', 'No summary available')
        digest.setdefault('event_calendar', [])
        digest.setdefault('action_items', [])
        digest.setdefault('important_announcements', [])
        
        # ---------------- Executive Summary ----------------
        self.doc.add_heading("Executive Summary", level=2)
        self.doc.add_paragraph(digest.get("executive_summary", "No summary available"))
        self.doc.add_paragraph("")

        # ---------------- Upcoming Events ----------------
        self.doc.add_heading("Upcoming Events", level=2)
        events = digest.get("event_calendar", [])
        
        # Ensure events is a list
        if not isinstance(events, list):
            logger.warning(f"event_calendar is not a list: {type(events)}")
            events = []
        
        logger.info(f"Rendering {len(events)} events in calendar")
            
        if not events:
            self.doc.add_paragraph("No events found.")
            logger.warning("No events to display in calendar")
        else:
            logger.info(f"Creating events table with {len(events)} rows")
            # Create table with headers
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Set column widths for better formatting
            # Event: 35%, Date: 25%, Time: 20%, Location: 20%
            table.columns[0].width = Inches(2.5)  # Event
            table.columns[1].width = Inches(1.8)  # Date
            table.columns[2].width = Inches(1.2)  # Time
            table.columns[3].width = Inches(1.5)  # Location
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Event'
            header_cells[1].text = 'Date'
            header_cells[2].text = 'Time'
            header_cells[3].text = 'Location'
            
            # Make headers bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add event rows
            for event in events:
                # Ensure event is a dict
                if not isinstance(event, dict):
                    logger.warning(f"Event is not a dict: {type(event)}")
                    continue
                    
                row_cells = table.add_row().cells
                row_cells[0].text = str(event.get('title', 'Event'))
                row_cells[1].text = str(event.get('date', 'TBD'))
                row_cells[2].text = str(event.get('time', ''))
                row_cells[3].text = str(event.get('location', ''))
            
            self.doc.add_paragraph("")  # Spacing after table
            
            # Add detailed event information below table
            self.doc.add_heading("Event Details", level=3)
            for event in events:
                # Ensure event is a dict
                if not isinstance(event, dict):
                    continue
                    
                title = str(event.get('title', 'Event'))
                details = str(event.get('details', ''))
                
                if details:
                    # Event name as bold subheading
                    p = self.doc.add_paragraph()
                    run = p.add_run(f"{title}:")
                    run.bold = True
                    
                    # Details
                    self.doc.add_paragraph(f"  {details}")
                    
                    # Sources
                    sources = event.get("sources", [])
                    if sources:
                        source_text = f"  (Mentioned in: {', '.join(sources)})"
                        source_p = self.doc.add_paragraph(source_text)
                        source_p.runs[0].font.size = Pt(9)
                    
                    self.doc.add_paragraph("")  # Spacing between events

        self.doc.add_paragraph("")

        # ---------------- Action Items ----------------
        self.doc.add_heading("Action Items", level=2)
        items = digest.get("action_items", [])
        
        # Ensure items is a list
        if not isinstance(items, list):
            logger.warning(f"action_items is not a list: {type(items)}")
            items = []
            
        if not items:
            self.doc.add_paragraph("No action items found.")
        else:
            for item in items:
                # Ensure item is a dict
                if not isinstance(item, dict):
                    logger.warning(f"Action item is not a dict: {type(item)}")
                    continue
                    
                priority = str(item.get('priority', 'medium')).upper()
                action = str(item.get('action', 'No action'))
                
                # Priority indicator
                priority_emoji = {'HIGH': '🔴', 'MEDIUM': '🟡', 'LOW': '🟢'}.get(priority, '⚪')
                
                text = f"{priority_emoji} [{priority}] {action}"
                
                # Add due date if present
                due_date = item.get('due_date')
                if due_date:
                    text += f" (Due: {due_date})"
                
                self.doc.add_paragraph(text)
                
                # Add details if present
                details = item.get('details', '')
                if details:
                    detail_p = self.doc.add_paragraph(f"  {str(details)}")
                    detail_p.runs[0].font.size = Pt(10)
                    
        self.doc.add_paragraph("")

        # ---------------- Announcements ----------------
        self.doc.add_heading("Important Announcements", level=2)
        anns = digest.get("important_announcements", [])
        
        # Ensure announcements is a list
        if not isinstance(anns, list):
            logger.warning(f"important_announcements is not a list: {type(anns)}")
            anns = []
            
        if not anns:
            self.doc.add_paragraph("None.")
        else:
            for ann in anns:
                # Handle both string and dict formats
                if isinstance(ann, dict):
                    ann_text = ann.get('announcement', str(ann))
                elif isinstance(ann, str):
                    ann_text = ann
                else:
                    ann_text = str(ann)
                    
                self.doc.add_paragraph(f"• {ann_text}")
        self.doc.add_page_break()

    def _add_email_summaries(self, emails):
        """Add section for individual email summaries."""
        self.doc.add_heading("Individual Email Summaries", level=2)

        if not emails:
            self.doc.add_paragraph("No emails summarized.")
            return

        for email in emails:
            # Email dict structure from orchestrator:
            # - email['summary'] is the full summary dict from AI summarizer
            # - Top level might also have subject/sender/date
            
            # Get metadata - check both top level and nested in summary dict
            subject = email.get("subject", "No subject")
            sender = email.get("sender", "Unknown sender")  
            date = email.get("date") or email.get("received", "Unknown date")

            self.doc.add_heading(subject, level=3)
            self.doc.add_paragraph(f"From: {sender} ({date})")

            # Extract the actual summary text
            summary_text = None
            summary_data = email.get("summary")
            
            if isinstance(summary_data, dict):
                # Case 1: summary is a dict from AI with nested 'summary' field
                summary_text = summary_data.get('summary', '')
                logger.debug(f"Extracted summary from dict for {subject}: {len(summary_text) if summary_text else 0} chars")
            elif isinstance(summary_data, str):
                # Case 2: summary is already a string
                summary_text = summary_data
                logger.debug(f"Using string summary for {subject}: {len(summary_text)} chars")
            
            # Fallback if no summary found
            if not summary_text:
                summary_text = "No summary available for this email."
                logger.warning(f"No summary found for email: {subject} (message_id: {email.get('message_id', 'unknown')})")
            
            # Clean up markdown formatting if present
            summary_text = self._clean_markdown(summary_text)

            # Split long text into paragraphs to avoid truncation
            # python-docx has issues with very long single paragraphs
            if len(summary_text) > 1000:
                # Split by double newlines first (natural paragraph breaks)
                paragraphs = summary_text.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        # If still too long, split by single newlines
                        if len(para_text) > 1000:
                            lines = para_text.split('\n')
                            for line in lines:
                                if line.strip():
                                    self.doc.add_paragraph(line.strip())
                        else:
                            self.doc.add_paragraph(para_text.strip())
            else:
                self.doc.add_paragraph(summary_text)
            
            self.doc.add_paragraph("")  # Spacing between emails
    
    def _clean_markdown(self, text: str) -> str:
        """Remove markdown formatting from text for clean Word doc display."""
        import re
        
        # Remove markdown headers (##, ###, etc.)
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        
        # Remove bold/italic markers (**, *, __, _)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # **bold**
        text = re.sub(r'\*([^*]+)\*', r'\1', text)      # *italic*
        text = re.sub(r'__([^_]+)__', r'\1', text)      # __bold__
        text = re.sub(r'_([^_]+)_', r'\1', text)        # _italic_
        
        # Remove bullet points at line start (-, *, +)
        text = re.sub(r'^\s*[-*+]\s+', '  • ', text, flags=re.MULTILINE)
        
        # Clean up excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()

    # ======================================================================
    # Text digest for email body
    # ======================================================================
    
    def create_simple_text_digest(
        self,
        digest_data: dict,
        email_summaries: list,
        date_str: str
    ) -> str:
        """Create a plain text version of the digest for email body.
        
        Args:
            digest_data: The consolidated digest dictionary
            email_summaries: List of email summary dictionaries
            date_str: Date string for the digest
            
        Returns:
            Plain text formatted digest
        """
        lines = []
        lines.append(f"CUSD Email Digest - {date_str}")
        lines.append("=" * 60)
        lines.append("")
        
        # Executive Summary
        lines.append("EXECUTIVE SUMMARY:")
        lines.append(digest_data.get("executive_summary", "No summary available"))
        lines.append("")
        
        # Events
        lines.append("UPCOMING EVENTS:")
        events = digest_data.get("event_calendar", [])
        if not events:
            lines.append("  No events found.")
        else:
            for event in events:
                event_line = f"  • {event.get('date', 'N/A')} - {event.get('event', 'No title')}"
                if event.get("time"):
                    event_line += f" at {event['time']}"
                if event.get("location"):
                    event_line += f" ({event['location']})"
                lines.append(event_line)
                
                # Add ELC note if Early Release
                if "Early Release" in event.get("event", ""):
                    lines.append("    Note: Student attends ELC - no pickup change needed.")
        lines.append("")
        
        # Action Items
        lines.append("ACTION ITEMS:")
        items = digest_data.get("action_items", [])
        if not items:
            lines.append("  No action items found.")
        else:
            for item in items:
                priority = item.get('priority', 'medium').upper()
                action = item.get('action', 'No action')
                item_line = f"  [{priority}] {action}"
                if item.get("due_date"):
                    item_line += f" (Due {item['due_date']})"
                lines.append(item_line)
        lines.append("")
        
        # Announcements
        lines.append("IMPORTANT ANNOUNCEMENTS:")
        anns = digest_data.get("important_announcements", [])
        if not anns:
            lines.append("  None.")
        else:
            for ann in anns:
                lines.append(f"  • {ann}")
        lines.append("")
        
        lines.append("=" * 60)
        lines.append(f"Total emails processed: {len(email_summaries)}")
        
        return "\n".join(lines)
